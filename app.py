import streamlit as st
import fitz  # PyMuPDF
import docx
import os
import zipfile
from tempfile import TemporaryDirectory

def extract_legal_rules(text):
    # هذه دالة تجريبية لاستخراج القواعد القضائية من النص
    # في النسخة النهائية يجب استبدالها بخوارزمية حقيقية
    lines = text.split("\n")
    rules = [line.strip() for line in lines if "قاعدة" in line or "حكم" in line]
    return rules if rules else ["لم يتم العثور على قواعد قضائية واضحة في هذا الحكم."]

def create_word_file(rules, filename):
    doc = docx.Document()
    doc.add_heading('القواعد القضائية المستخرجة', level=1)
    for i, rule in enumerate(rules, 1):
        doc.add_paragraph(f"{i}- {rule}", style='List Number')
    doc.save(filename)

st.set_page_config(page_title="استخراج القواعد القضائية من ملفات PDF", layout="centered")
st.title("📄 استخراج القواعد القضائية من عدة أحكام PDF")

uploaded_files = st.file_uploader("📂 قم برفع ملفات الأحكام بصيغة PDF", type="pdf", accept_multiple_files=True)

if uploaded_files:
    if st.button("🔍 استخراج القواعد وإنشاء ملف ZIP"):
        with TemporaryDirectory() as temp_dir:
            zip_path = os.path.join(temp_dir, "قواعد_قضائية_لكل_حكم.zip")
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for uploaded_file in uploaded_files:
                    pdf_name = uploaded_file.name.replace('.pdf', '')
                    docx_path = os.path.join(temp_dir, f"{pdf_name}.docx")

                    # قراءة النص من PDF
                    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                    full_text = "\n".join([page.get_text() for page in doc])
                    doc.close()

                    # استخراج القواعد وكتابة ملف وورد
                    rules = extract_legal_rules(full_text)
                    create_word_file(rules, docx_path)

                    # إضافة إلى الملف المضغوط
                    zipf.write(docx_path, arcname=os.path.basename(docx_path))

            with open(zip_path, "rb") as f:
                st.download_button(
                    label="📥 تحميل الملف المضغوط للقواعد القضائية",
                    data=f,
                    file_name="قواعد_قضائية_لكل_حكم.zip",
                    mime="application/zip"
                )
